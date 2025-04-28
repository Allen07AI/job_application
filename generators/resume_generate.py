import docx
import openai
import asyncio
from agents import Agent, Runner, OpenAIChatCompletionsModel, AsyncOpenAI
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.section import WD_SECTION_START
from dotenv import load_dotenv
import os

# Load environment variables and set OpenAI API key
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
client = openai.OpenAI()

# --- Helper functions ---

# Define the Resume Tailoring Agent
resume_tailoring_agent = Agent(
    name="Resume Tailoring Agent",
    instructions=(
        "Your task is to modify the given resume section to match the keywords from the job description. "
        "Keep all relevant information intact while optimizing the phrasing for applicant tracking systems (ATS). "
        "Ensure each bullet point remains concise and action-driven. Always limit it to 3 points. Never do more points than 3."
        "Do not make up stuff and keep the resume poins intact but just modify according to the keywords in the job description."
    ),
    model=OpenAIChatCompletionsModel(
        model="gpt-3.5-turbo",
        openai_client= AsyncOpenAI()
    ),
)

# Function to tailor resume synchronously
def tailor_section(job_desc, resume_section, max_points):
    """Uses the resume tailoring agent to refine the resume section based on the job description."""
    input_text = (
        f"Job Description:\n{job_desc}\n\n"
        f"Resume Section to Tailor (generate at most {max_points} bullet points):\n{resume_section}\n\n"
        "Modify the resume section to align with the job description. Keep experience and domain unchanged."
        "Do not include the title of the job or working dates or position. Your job is to only edit the points."
        "Make sure to add a hyphen (-) as a bullet point for each point."
    )

    result = asyncio.run(Runner.run(resume_tailoring_agent, input=input_text))
    return result.final_output

def apply_standard_style(paragraph):
    """Set the paragraph font to Times New Roman, size 12, and line spacing 1"""
    paragraph.style.font.name = 'Times New Roman'
    for run in paragraph.runs:
        run.font.size = Pt(11)
    paragraph.paragraph_format.line_spacing = 1

def format_bullet_points(text):
    """Formats bullet points by ensuring they start with '• ' and removes any dashes after the bullet."""
    lines = text.split("\n")
    formatted_lines = []
    for line in lines:
        line = line.strip()
        if line.startswith("-"):
            formatted_lines.append("• " + line[1:].strip())  
        elif line.startswith("•"):
            formatted_lines.append(line) 
        else:
            formatted_lines.append(line)  
    return "\n".join(formatted_lines)
   
    
def add_work_proj_text(doc, text):
    """Adds a normal text paragraph with standard style."""
    para = doc.add_paragraph(format_bullet_points(text))
    apply_standard_style(para)
    return para

def add_heading(doc, text):
    """Adds a bold heading with an extended underline effect."""
    heading_para = doc.add_paragraph()
    run = heading_para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.underline = True
    heading_para.paragraph_format.space_after = Pt(0.5)

def add_body_text(doc, text):
    """Adds a normal text paragraph with standard style."""
    para = doc.add_paragraph(text)
    apply_standard_style(para)
    return para

def add_sub_heading(doc, text):
    """Adds a bolded text paragraph with standard style."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    apply_standard_style(para)
    return para

def add_italic_heading(doc, text):
    """Adds a bolded text paragraph with standard style."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    para.paragraph_format.line_spacing = 1
    return para


def tailor_section_gpt(job_desc, section_text, max_points):
    #Calls GPT to generate tailored bullet points. Limits output to max_points bullet points.
    prompt = (
        f"Job Description:\n{job_desc}\n\n"
        f"Resume Section to Tailor (generate at most {max_points} bullet points):\n{section_text}\n\n"
        "Modified Section (only include relevant bullet points based solely on keywords from the job description):\n"
        "Don't drift away from my work experience and project input. Just tailor the input to the keywords from the Job Description. Always add a bullet point before each point. Keep all the relevant information intact.\n"
        "Do not change the domain of the projects or work experience. The keywords are mostly action verbs"
    )
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are an expert resume writer that tailors work experience and projects to match the keywords in the job description."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=200,
        temperature=0.1
    )
    return response.choices[0].message.content.strip()

# --- Main resume update function ---

def generate_resume(job_desc, output_path, user_data):
    doc = docx.Document()

    # Set page margins (0.5 cm top and bottom, 1 cm left and right)
    section = doc.sections[0]
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    # Preventing additional pages
    section.start_type = WD_SECTION_START.NEW_PAGE

    # Name (Centered, Bold, Underlined)
    title_para = doc.add_paragraph()
    run = title_para.add_run(user_data["name"])
    run.bold = True
    run.underline = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.line_spacing = 1

    # Contact Info (Centered)
    contact_para = doc.add_paragraph(f"Phone: {user_data['contact_info_num']} | Email: {user_data['contact_info_email']}")
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    apply_standard_style(contact_para)

    # Education
    add_heading(doc, "EDUCATION" + " " + "_" * 88)
    add_sub_heading(doc, f"{user_data['university']} - GPA: {user_data['gpa']}")
    add_body_text(doc, user_data["course"])
    add_body_text(doc, f"Relevant Coursework: {user_data['relevant_coursework']}")

    # Skills
    add_heading(doc, "SKILLS" + " " + "_" * 93)
    add_body_text(doc, f"Software/Tools: {user_data['software_tools']}")
    add_body_text(doc, f"Interests: {user_data['interests']}")

    # Work Experience
    add_heading(doc, "WORK EXPERIENCE" + " " + "_" * 80)
    
    work_ex_1 = f"{user_data['company_1']} ({user_data['dates_1']})\n{user_data['work_1']}"
    tailored_work_ex_1 = tailor_section(job_desc, work_ex_1, 3)
    add_sub_heading(doc, f"{user_data['company_1']} | {user_data['dates_1']}")
    add_italic_heading(doc, user_data["company_1_job_title"])
    add_work_proj_text(doc, tailored_work_ex_1)

    work_ex_2 = f"{user_data['company_2']} ({user_data['dates_2']})\n{user_data['work_2']}"
    tailored_work_ex_2 = tailor_section(job_desc, work_ex_2, 3)
    add_sub_heading(doc, f"{user_data['company_2']} | {user_data['dates_2']}")
    add_italic_heading(doc, user_data["company_2_job_title"])
    add_work_proj_text(doc, tailored_work_ex_2)

    # Projects
    add_heading(doc, "PROJECTS" + " " + "_" * 90)

    project_1 = f"{user_data['project_1_title']}\n{user_data['project_1_desc']}"
    tailored_project_1 = tailor_section(job_desc, project_1, 3)
    add_sub_heading(doc, user_data['project_1_title'])
    add_work_proj_text(doc, tailored_project_1)

    project_2 = f"{user_data['project_2_title']}\n{user_data['project_2_desc']}"
    tailored_project_2 = tailor_section(job_desc, project_2, 3)
    add_sub_heading(doc, user_data['project_2_title'])
    add_work_proj_text(doc, tailored_project_2)

    # Save the document
    doc.save(output_path)
    print(f"Updated resume saved as {output_path}")
