import docx
import openai
from openai import OpenAI
import asyncio
from agents import Agent, Runner, OpenAIChatCompletionsModel, AsyncOpenAI
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
import ai_agents
from ai_agents import generate_cover_letter_content
from dotenv import load_dotenv
import os
from datetime import datetime
import pandas as pd

# OpenAI GPT model for generating cover letter content
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
client = openai.OpenAI()

def generate_cover_letter_gpt(job_title, company_name, user_data):
    prompt = f"""Write a professional cover letter for a {job_title} position at {company_name}. Keep it to a maximum of two paragraphs, and include enthusiasm for the role and relevant skills. 
    Don't generate fake experience or add irrelevant stuff. Generate only the body of the Cover Letter. Do not generate anything like formats or subjects. 

    Start with 'I am excited to apply to'. Mention that I am currently pursuing a {user_data['course']} at {user_data['university']}. Highlight my experience with {', '.join(user_data['software_tools'])}, and how it applies to the role. 
    Briefly reference my previous role as a {user_data['company_1_job_title']} at {user_data['company_1']}, where I worked on {user_data['work_1']}. 
    Also, briefly reference my previous role as a {user_data['company_2_job_title']} at {user_data['company_2']}, where I worked on {user_data['work_2']}.
    Ensure the response is concise, engaging, and directly relevant to the role.""" 
    messages = [
        {"role": "system", "content": "You are a helpful assistant who writes the body paragraphs for professional and concise cover letters."},
        {"role": "user", "content": prompt}
    ]
    response = client.chat.completions.create(
        model="gpt-3.5-turbo", 
        messages = messages, 
        max_tokens=200, 
        temperature= 0.1)
    return response.choices[0].message.content.strip()
    

def create_cover_letter(user_data, job_title, company_name, save_dir):
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    
    # Set font style
    def set_font(paragraph):
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # Set line spacing to 1.5
    def set_line_spacing(paragraph):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.5

    # Cover Letter Header (Bold and Underlined)
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    run = header.add_run("Cover Letter")
    run.bold = True
    run.underline = True 
    set_font(header)
    set_line_spacing(header)

    # Add user information
    doc.add_paragraph(user_data["name"])
    doc.add_paragraph(user_data["contact_info_num"])
    doc.add_paragraph(datetime.today().strftime("%m/%d/%Y"))

    # Add company details and subject
    doc.add_paragraph(f"\nHiring Committee\n{company_name}")
    doc.add_paragraph(f"Subject: Application for {job_title}")

    # Add greeting
    doc.add_paragraph("\nDear Hiring Committee,")

    # Generate AI content for the body of the letter
    cover_letter_body = generate_cover_letter_content(job_title, company_name, user_data)
    doc.add_paragraph(cover_letter_body)

    # Closing
    doc.add_paragraph("\nSincerely,")
    doc.add_paragraph(user_data["name"])

    # Apply font style and line spacing to the whole document
    for para in doc.paragraphs:
        set_font(para)
        set_line_spacing(para)

    # Save the document
    doc.save(save_dir)
    print(f"Updated resume saved as {save_dir}")
   